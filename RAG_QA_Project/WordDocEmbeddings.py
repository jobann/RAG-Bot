from docx import Document
from chromadb import PersistentClient
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
import json


class WordDocEmbeddings:
    def __init__(self, docx_path: str, persist_dir: str = "./chroma_word_rich", collection_name: str = "word_full_format"):
        self.docx_path = docx_path
        self.blocks = []
        self.block_id = 0

        self.client = PersistentClient(path=persist_dir)
        self.embedding_function = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")

        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            embedding_function=self.embedding_function
        )

    def _get_font_properties(self, run, paragraph):
        font = run.font
        fallback_font = paragraph.style.font

        props = {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "strikethrough": font.strike,
            "subscript": font.subscript,
            "superscript": font.superscript,
            "font_name": font.name or fallback_font.name,
            "font_size": font.size.pt if font.size else (
                fallback_font.size.pt if fallback_font.size else None
            ),
            "text_color": str(font.color.rgb) if font.color and font.color.rgb else None,
            "highlight_color": font.highlight_color.name if font.highlight_color else None,
            "all_caps": font.all_caps,
            "small_caps": font.small_caps,
            "shadow": font.shadow,
            "outline": font.outline,
            "emboss": font.emboss,
            "engrave": font.imprint
        }

        return {k: v for k, v in props.items() if v is not None}

    def _get_paragraph_formatting(self, paragraph):
        fmt = paragraph.paragraph_format
        props = {
            "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
            "line_spacing": fmt.line_spacing,
            "space_before": fmt.space_before.pt if fmt.space_before else None,
            "space_after": fmt.space_after.pt if fmt.space_after else None,
            "left_indent": fmt.left_indent.pt if fmt.left_indent else None,
            "right_indent": fmt.right_indent.pt if fmt.right_indent else None,
            "first_line_indent": fmt.first_line_indent.pt if fmt.first_line_indent else None,
            "keep_together": fmt.keep_together,
            "keep_with_next": fmt.keep_with_next,
            "widow_control": fmt.widow_control,
            "style": paragraph.style.name
        }
        
        return {k: v for k, v in props.items() if v is not None}


    def _process_paragraphs(self, paragraphs):
        for para in paragraphs:
            if not para.text.strip():
                continue

            para_meta = self._get_paragraph_formatting(para)
            runs_meta = [
                {
                    "text": run.text,
                    **self._get_font_properties(run, para)
                }
                for run in para.runs if run.text.strip() and self._get_font_properties(run, para)
            ]

            self.blocks.append({
                "id": f"block_{self.block_id}",
                "text": para.text.strip(),
                "metadata": {
                    "type": "paragraph",
                    "paragraph_format": para_meta,
                    "runs": runs_meta
                }
            })
            self.block_id += 1

    def _extract_cell_metadata(self, cell):
        cell_paragraphs = cell.paragraphs
        para_data = []

        final_metadata = {
            "type": "table_cell",
            "style": cell.paragraphs[0].style.name if cell.paragraphs else None
        }

        for para in cell_paragraphs:
            if not para.text.strip():
                continue
            runs = {}
            for run in para.runs:
                if run.text.strip():
                    font_props = self._get_font_properties(run, para)
                    if font_props:  # only append if there's at least one property
                        final_metadata["font_props"] = font_props

        return final_metadata

    def _process_tables(self, tables):
        for t_index, table in enumerate(tables):
            table_info = []
            for r_index, row in enumerate(table.rows):
                for c_index, cell in enumerate(row.cells):
                    table_info.append({
                        "text": cell.text.strip(),
                        "row_index": r_index,
                        "col_index": c_index,
                        "metadata": self._extract_cell_metadata(cell)
                    })

            self.blocks.append({
                "id": f"block_{self.block_id}",
                "table_index": t_index,
                "table_info": table_info
            })
            self.block_id += 1

    def extract(self):
        print(f"Loading Word document from: {self.docx_path}")
        doc = Document(self.docx_path)
        self._process_paragraphs(doc.paragraphs)
        self._process_tables(doc.tables)
        print(f"Extracted {len(self.blocks)} formatted blocks.")

    def store(self):
        print(f"Storing into ChromaDB...")
        for block in self.blocks:
            self.collection.add(
                ids=[block["id"]],
                documents=[block["text"]],
                metadatas=[block["metadata"]]
            )
        print("Storage complete.")

    def print_blocks(self):
        for block in self.blocks:
            print(json.dumps(block, indent=2))
            # print(f"ID: {block['id']}")
            # if "text" in block:
                # print(f"Text: {block['text']}")
                # print(json.dumps(block, indent=2))
            # elif "table_index" in block:
                # print("Table Info:")
                # print(block)
                # print(json.dumps(block, indent=2))
                # for cell in block["table_info"]:
                #     print(f"  ({cell['row_index']}, {cell['col_index']}) ➜ {cell['text']}")
            # else:
                # print("No displayable text found.")
            # print("Metadata:")
            # metadata = block.get("metadata") or {k: v for k, v in block.items() if k not in {"id", "text", "table_info"}}
            # print(json.dumps(metadata, indent=2))
            print("=" * 50)

    def run(self):
        self.extract()
        print("\n\n\n")
        self.print_blocks()
        # self.store()



if __name__ == "__main__":
    wordDocEmbeddings = WordDocEmbeddings(docx_path="./data/Sample_WordDocument.docx")
    wordDocEmbeddings.run()