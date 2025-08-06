from docx import Document
from chromadb import PersistentClient
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
import json
import requests


class WordDocEmbeddings:
    def __init__(self, docx_path: str, persist_dir: str = "./chroma_word_rich_NL", collection_name: str = "word_full_format"):
        self.docx_path = docx_path
        self.blocks = []
        self.block_id = 0

        self.client = PersistentClient(path=persist_dir)
        self.embedding_function = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")

        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            embedding_function=self.embedding_function
        )
    

    def _get_font_properties_natural_language(self, run, paragraph):
        font = run.font
        fallback_font = paragraph.style.font

        formatted = []

        if run.bold: formatted.append("bold")
        if run.italic: formatted.append("italic")
        if run.underline: formatted.append("underlined")
        if font.strike: formatted.append("strikethrough")
        if font.subscript: formatted.append("subscript")
        if font.superscript: formatted.append("superscript")
        if font.name:
            formatted.append(f"font {font.name}")
        elif fallback_font.name:
            formatted.append(f"font {fallback_font.name}")
        if font.size:
            formatted.append(f"with font size {font.size.pt}pt")
        elif fallback_font.size:
            formatted.append(f"{fallback_font.size.pt}pt")
        if font.color and font.color.rgb:
            formatted.append(f"colored #{str(font.color.rgb)}")
        if font.highlight_color:
            formatted.append(f"highlighted in {font.highlight_color.name}")
        if font.all_caps:
            formatted.append("all caps")
        if font.small_caps:
            formatted.append("small caps")
        if font.shadow:
            formatted.append("shadowed")
        if font.outline:
            formatted.append("outlined")
        if font.emboss:
            formatted.append("embossed")
        if font.imprint:
            formatted.append("engraved")

        return ", ".join(formatted) if formatted else None


    def _get_paragraph_formatting_natural_language(self, paragraph):
        fmt = paragraph.paragraph_format
        descriptions = []

        if paragraph.alignment is not None:
            descriptions.append(f"alignment is {paragraph.alignment}")
        if fmt.line_spacing:
            descriptions.append(f"line spacing is {fmt.line_spacing}")
        if fmt.space_before:
            descriptions.append(f"space before is {fmt.space_before.pt}pt")
        if fmt.space_after:
            descriptions.append(f"space after is {fmt.space_after.pt}pt")
        if fmt.left_indent:
            descriptions.append(f"left indent is {fmt.left_indent.pt}pt")
        if fmt.right_indent:
            descriptions.append(f"right indent is {fmt.right_indent.pt}pt")
        if fmt.first_line_indent:
            descriptions.append(f"first line indent is {fmt.first_line_indent.pt}pt")
        if fmt.keep_together:
            descriptions.append("keep lines together is enabled")
        if fmt.keep_with_next:
            descriptions.append("keep with next is enabled")
        if fmt.widow_control:
            descriptions.append("widow control is enabled")
        if paragraph.style and paragraph.style.name and paragraph.style.name.strip().lower() != "normal":
            descriptions.append(f"paragraph style is '{paragraph.style.name}'")

        return ". ".join(descriptions) + "." if descriptions else None

    def _process_paragraphs(self, paragraphs):
        for para in paragraphs:
            if not para.text.strip():
                continue

            para_meta = self._get_paragraph_formatting_natural_language(para)
            inline_text_parts = []

            for run in para.runs:
                run_text = run.text.strip()
                if not run_text:
                    continue
                font_desc = self._get_font_properties_natural_language(run, para)
                if font_desc:
                    inline_text_parts.append(f"~!{run_text}~!({font_desc})")
                else:
                    inline_text_parts.append(run_text)

            full_text = " ".join(inline_text_parts)

            if para_meta and para_meta.strip().lower() != "paragraph style is 'normal'.":
                full_text += f" ({para_meta})"

            self.blocks.append({
                "id": f"block_{self.block_id}",
                "text": full_text
            })

            self.block_id += 1


    def _extract_cell_metadata_natural_language(self, cell):
        cell_paragraphs = cell.paragraphs
        descriptions = []

        for para in cell_paragraphs:
            if not para.text.strip():
                continue
            for run in para.runs:
                if run.text.strip():
                    font_props = self._get_font_properties_natural_language(run, para)
                    if font_props:
                        descriptions.append(f'{font_props}')

        return " ".join(descriptions) if descriptions else None

 

    def _process_tables_natural_language(self, tables):
        for t_index, table in enumerate(tables):
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            summary_lines = []

            for r_index, row in enumerate(table.rows[1:], start=1):
                cell_sentences = []
                for c_index, cell in enumerate(row.cells):
                    if c_index >= len(headers):
                        continue
                    header = headers[c_index]
                    if not header:
                        continue
                    cell_text = cell.text.strip() or "empty"
                    cell_font_desc = self._extract_cell_metadata_natural_language(cell)
                    para_format_desc = self._get_paragraph_formatting_natural_language(cell.paragraphs[0])
                    descriptions = [d for d in [cell_font_desc, para_format_desc] if d]
                    formatting_desc = " | ".join(descriptions) if descriptions else None
                    if formatting_desc:
                        cell_text = f"~!{cell_text}~!({formatting_desc})"
                    cell_sentences.append(f"{cell_text} is the {header}")
                if cell_sentences:
                    summary_lines.append(", ".join(cell_sentences) + ".")

            summary = " ".join(summary_lines)

            self.blocks.append({
                "id": f"block_{self.block_id}",
                "text": summary,
                "metadata": {
                    "table_index": t_index
                }
            })
            self.block_id += 1

    


    def _generate_embedding_input(self, block):
        if "text" in block:  # Paragraph block
            base_text = block["text"]
            metadata = block.get("metadata")
            if metadata:
                metadata_str = json.dumps(metadata, indent=2)
                return f"{base_text}\n\nMetadata:\n{metadata_str}"
            return base_text

        elif "table_info" in block:  # Table block
            table_rows = block["table_info"]["rows"]
            if not table_rows or len(table_rows[0]) < 2:
                return "Table data insufficient"
            header = [cell["text"] for cell in table_rows[0]]
            table_text_lines = []
            for row in table_rows[1:]:  # skip header row
                if len(row) >= len(header):
                    sentence_parts = [f"{header[i]} is {row[i]['text']}" for i in range(len(header))]
                    table_text_lines.append(". ".join(sentence_parts) + ".")
            return "\n".join(table_text_lines)

        else:
            return "Unknown block structure"

    def _flatten_metadata(self, metadata: dict, parent_key: str = "", sep: str = "_"):
        flat = {}
        for k, v in metadata.items():
            key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                flat.update(self._flatten_metadata(v, key, sep=sep))
            elif isinstance(v, list):
                # Convert lists to JSON strings to preserve structure in string format
                flat[key] = json.dumps(v)
            else:
                flat[key] = v
        return flat

    def store(self):
        print(f"Storing into ChromaDB...")
        for block in self.blocks:
            embedding_input = self._generate_embedding_input(block)
            flat_metadata = self._flatten_metadata(block.get("metadata", {}))

            if flat_metadata:
                self.collection.add(
                    ids=[block["id"]],
                    documents=[embedding_input],
                    metadatas=[flat_metadata]
                )
            else:
                self.collection.add(
                    ids=[block["id"]],
                    documents=[embedding_input]
                )
        print("Storage complete.")




    def ask_llama(self, query: str, top_k: int = 5, model: str = "llama3.2"):
        print(f"\n\n{query}")

        results = self.collection.query(
            query_texts=[query],
            n_results=top_k
        )


        context = "\n\n".join(results["documents"][0])


        prompt = f"""Question: {query}

            Instructions for Answering:

                Act like a professional assistant.

                ~ Answer naturally and concisely, using only the information provided in the Context.

                ~ text within ~! marker up to the next ~! marker is formatted and after the second ~! marker style is described in parentheses, including punctuation or connecting words (e.g., parentheses, commas, "and", etc.), as part of the styled content. for example ~!CIBC) and~!(bold, italic), here 'CIBC) and' is formatted as both bold and italic, dont reflect those in your response until ask specifically in question.

                ~ Do not speculate or add information that is not explicitly in the Context.

                ~ If multiple items are relevant, include all.

                ~ If the answer is negative or not found in the Context, say so clearly and politely.

            Context: {context}"""
        
        print(prompt)
        
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False
            }
        )

        if response.status_code == 200:
            answer = response.json()["response"]
            import re
            # Remove ~! markers and anything in parentheses after it
            cleaned_answer = re.sub(r'~!(.*?)~!\(.*?\)', r'\1', answer)
            print(f"LLaMA Answer:\n{cleaned_answer.strip()}\n\n")
        else:
            print(f"Failed to query LLaMA: {response.status_code}, {response.text}")
    
    
    def extract(self):
            print(f"Loading Word document from: {self.docx_path}")
            doc = Document(self.docx_path)
            self._process_paragraphs(doc.paragraphs)
            self._process_tables_natural_language(doc.tables)
            print(f"Extracted {len(self.blocks)} formatted blocks.")


    def print_blocks(self):
        for block in self.blocks:
            print(json.dumps(block, indent=2))
            print("=" * 50)
    

    def run(self):
        self.extract()
        print("\n\n\n")
        # self.print_blocks()
        self.store()



if __name__ == "__main__":
    wordDocEmbeddings = WordDocEmbeddings(docx_path="./data/Sample_WordDocument.docx")
    wordDocEmbeddings.run()
    # wordDocEmbeddings.ask_llama("Which bank is in Brampton?")
    wordDocEmbeddings.ask_llama("Which text is bold in document?")

    # Load questions and references from file
    # with open("./data/qa_pairs.json", "r", encoding="utf-8") as f:
    #     all_qa_pairs = json.load(f)

    # question_indices_to_run = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10]  
    # # question_indices_to_run = [9]  

    # if question_indices_to_run:
    #     qa_pairs = [all_qa_pairs[i] for i in question_indices_to_run if i < len(all_qa_pairs)]
    # else:
    #     qa_pairs = all_qa_pairs


    # for qa in qa_pairs:
    #     print(f"\n{'='*80}")
    #     wordDocEmbeddings.ask_llama(qa['question'])